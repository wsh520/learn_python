from docx import Document
import os
def docx_to_markdown(docx_file):
    doc = Document(docx_file)
    markdown_content = []
    images = []

    # 遍历文档中的每个段落
    for para in doc.paragraphs:
        # 处理标题
        if para.style.name.startswith('Heading'):
            level = len(para.style.name) - len(para.style.name.rstrip('0123456789'))
            text = para.text.strip()
            markdown_content.append('#' * level + ' ' + text)
        # 处理普通文本
        else:
            text = para.text.strip()
            markdown_content.append(text)

    # 处理图像
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part.blob
            image_filename = rel.target_part.partname
            image_extension = os.path.splitext(image_filename)[1]
            image_filename = "image" + image_extension
            images.append((image_filename, image))
            # 在文本中插入图像的占位符
            markdown_content.append(f'![Image]({image_filename})\n')

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            row_content = '|'
            for cell in row.cells:
                row_content += cell.text.strip() + '|'
            markdown_content.append(row_content + '\n')
        markdown_content.append('|' + '---|' * len(table.columns) + '\n')

    # 将Markdown内容写入文件
    with open('output.md', 'w', encoding='utf-8') as md_file:
        for line in markdown_content:
            md_file.write(line + '\n')

    # 保存图像
    for image_filename, image in images:
        with open(image_filename, 'wb') as img_file:
            img_file.write(image)

    print("转换完成，Markdown文件已保存为output.md，图像已保存在当前目录。")

# 使用示例
docx_file_path = 'D:\item\python\mylearn\LABOMM相关bug和问题整理.docx'  # 替换为您的Word文档路径
docx_to_markdown(docx_file_path)